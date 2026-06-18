#!/usr/bin/env python3
"""
make_docx.py — 根据简单的 JSON 规格,为面试复盘/档案生成排好版的 .docx,
使用 python-docx(无需 LibreOffice)。

用法:
    python make_docx.py spec.json output.docx

规格格式见 references/output-templates.md
块类型:h1, h2, p(可选 "emphasis": true)、bullet(items[])、
       num(items[])、table(header[], rows[][], 可选 widths[] 单位 DXA)。
依赖:pip install python-docx --break-system-packages
"""
import json
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Microsoft YaHei"
NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x54, 0x96)
RED = RGBColor(0xC0, 0x00, 0x00)
HDR_FILL = "2E5496"
LABEL_FILL = "EAF0F8"


def set_cjk(run, font=FONT):
    run.font.name = font
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rfonts.set(qn(a), font)


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), fill)
    tcpr.append(sh)


def add_run(p, text, size=11, bold=False, color=None):
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    set_cjk(run)
    return run


def main():
    spec_path, out_path = sys.argv[1], sys.argv[2]
    with open(spec_path, encoding="utf-8") as f:
        spec = json.load(f)

    doc = Document()
    # base style
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(11)
    set_cjk(normal.element.rPr.rFonts if normal.element.rPr is not None else None) if False else None

    # page size US Letter
    sec = doc.sections[0]
    sec.page_width = Twips(12240)
    sec.page_height = Twips(15840)
    for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, m, Inches(1))

    # title
    if spec.get("title"):
        p = doc.add_paragraph()
        add_run(p, spec["title"], size=20, bold=True, color=NAVY)
    if spec.get("meta"):
        p = doc.add_paragraph()
        add_run(p, spec["meta"], size=10, color=RGBColor(0x88, 0x88, 0x88))

    for b in spec.get("blocks", []):
        t = b.get("type")
        if t == "h1":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            add_run(p, b["text"], size=15, bold=True, color=BLUE)
        elif t == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            add_run(p, b["text"], size=12.5, bold=True, color=NAVY)
        elif t == "p":
            p = doc.add_paragraph()
            emph = b.get("emphasis")
            add_run(p, b["text"], size=11, bold=bool(emph),
                    color=RED if emph else None)
        elif t == "bullet":
            for it in b["items"]:
                p = doc.add_paragraph(style="List Bullet")
                add_run(p, it, size=11)
        elif t == "num":
            for it in b["items"]:
                p = doc.add_paragraph(style="List Number")
                add_run(p, it, size=11)
        elif t == "table":
            header = b.get("header", [])
            rows = b.get("rows", [])
            ncol = len(header) if header else (len(rows[0]) if rows else 0)
            tbl = doc.add_table(rows=0, cols=ncol)
            tbl.style = "Table Grid"
            widths = b.get("widths")
            if header:
                hr = tbl.add_row().cells
                for i, htext in enumerate(header):
                    shade(hr[i], HDR_FILL)
                    pr = hr[i].paragraphs[0]
                    add_run(pr, htext, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            for row in rows:
                cells = tbl.add_row().cells
                for i, val in enumerate(row):
                    if i == 0 and not header:
                        shade(cells[i], LABEL_FILL)
                    pr = cells[i].paragraphs[0]
                    add_run(pr, str(val), size=10)
            if widths:
                for ri, r in enumerate(tbl.rows):
                    for i, c in enumerate(r.cells):
                        if i < len(widths):
                            c.width = Twips(widths[i])

    # Fix python-docx settings.xml schema nit: <w:zoom> needs w:percent
    try:
        settings_el = doc.settings.element
        zoom = settings_el.find(qn('w:zoom'))
        if zoom is not None and zoom.get(qn('w:percent')) is None:
            zoom.set(qn('w:percent'), '100')
    except Exception:
        pass

    doc.save(out_path)
    print("saved", out_path)


if __name__ == "__main__":
    main()
